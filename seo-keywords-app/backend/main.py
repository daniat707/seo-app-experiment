# backend/main.py
import os
import re
import json
import uuid
import tempfile
from typing import List, Dict

import pdfplumber
import docx
import pandas as pd
from docx import Document
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from openai import OpenAI
from pytrends.request import TrendReq


# ---------- FastAPI app & CORS ----------
app = FastAPI(title="SEO Keyword Finder")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173", "*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------- OpenAI client ----------
# Requires OPENAI_API_KEY in the environment
client = OpenAI()


# ---------- Helpers ----------
def extract_text_from_pdf(path: str) -> str:
    chunks = []
    with pdfplumber.open(path) as pdf:
        for p in pdf.pages:
            txt = p.extract_text() or ""
            if txt:
                chunks.append(txt)
    return "\n".join(chunks)


def extract_text_from_docx(path: str) -> str:
    d = docx.Document(path)
    return "\n".join([p.text for p in d.paragraphs if p.text])


def _clean_json_block(s: str) -> str:
    """
    Remove code fences and isolate the first { ... } JSON block.
    """
    s = s.strip()
    if s.startswith("```"):
        s = s.strip("`")
        s = "\n".join(s.splitlines()[1:])
    start = s.find("{")
    end = s.rfind("}")
    if start != -1 and end != -1 and end > start:
        return s[start : end + 1]
    return s


def generate_keywords_from_doc(doc_text: str) -> Dict:
    """
    Ask the model for English keywords + structure, even if the input is Spanish.
    """
    prompt = f"""
You are an SEO analyst. Analyze the document and return ONLY JSON (no prose).
Always output in ENGLISH, even if the source text is in Spanish.
Use widely searched English terms that a traveler from Europe/US/Canada/Asia would type.

Return EXACTLY this schema:
{{
  "language_detected": "<source language>",
  "primary_topics": ["..."],
  "seed_keywords": ["<15 short, high-signal head terms in ENGLISH>"],
  "long_tail_keywords": ["<20 intent-rich long-tail phrases in ENGLISH>"],
  "by_intent": {{
    "informational": ["..."],
    "commercial": ["..."],
    "transactional": ["..."],
    "navigational": ["..."]
  }},
  "questions": ["<10 common user queries in ENGLISH>"]
}}

Reference text (truncate if needed):
\"\"\"{doc_text[:12000]}\"\"\"
""".strip()

    response = client.responses.create(
        model="gpt-5",
        input=prompt,
        text={"verbosity": "low"}
    )
    output_text = getattr(response, "output_text", None) or str(response)
    payload = _clean_json_block(output_text)
    data = json.loads(payload)
    data.setdefault("seed_keywords", [])
    data.setdefault("long_tail_keywords", [])
    return data


def chunk(lst: List[str], n: int) -> List[List[str]]:
    return [lst[i : i + n] for i in range(0, len(lst), n)]


def trends_popularity(
    keywords: List[str], timeframe: str = "today 3-m", geo: str = ""
) -> Dict[str, float]:
    """
    Google Trends returns relative interest (0-100). We average over timeframe.
    Trends allows max 5 terms per request; we batch.
    """
    if not keywords:
        return {}

    pytrends = TrendReq(hl="en-US", tz=360)
    scores: Dict[str, float] = {}

    for group in chunk(keywords, 5):
        try:
            pytrends.build_payload(group, timeframe=timeframe, geo=geo)
            df = pytrends.interest_over_time()
            if df is None or df.empty:
                for k in group:
                    scores.setdefault(k, None)
                continue
            df = df.drop(columns=["isPartial"], errors="ignore")
            means = df.mean(numeric_only=True).to_dict()
            for k in group:
                scores[k] = float(means.get(k)) if k in means else None
        except Exception:
            for k in group:
                scores.setdefault(k, None)

    return scores


def compose_seo_copy_en(doc_text: str, ranked_keywords: List[str]) -> str:
    """
    Generate ORIGINAL English copy following the Spanish template, with keywords woven in.
    Returns Markdown; at the bottom include JSON-LD FAQPage.
    """
    top_kw = ", ".join(ranked_keywords[:15]) if ranked_keywords else ""
    prompt = f"""
You are a senior SEO copywriter for an inbound tour operator in Ecuador targeting travelers from
Europe, the US/Canada, and occasionally Asia. Write ORIGINAL ENGLISH copy (no plagiarism).

Follow this EXACT structure (H2/H3 + bullets):
## 1. General destination information
### Country overview (1–2 intro paragraphs)
### Reasons to visit (culture, nature, adventure, gastronomy)
### Highlight regions or zones (Ecuador: Andes/Highlands, Amazon, Coast, Galápagos; Peru: Coast, Andes, Amazon)
### Climate & best time to visit (clear, practical)
### Entry requirements (visas, vaccines) – include a disclaimer to verify official sources
### Map or infographic (short caption text describing what to show)

Rules:
- Weave these keywords naturally (no stuffing, ~1–2% density): {top_kw}
- Tone: clear, trustworthy, inspiring; emphasize authenticity, sustainability, safety.
- Avoid medical/legal claims; use soft guidance and official-source disclaimers.
- End with:
  - Meta title (≤ 60 chars)
  - Meta description (≤ 155 chars)
  - Suggested slugs (5)
  - JSON-LD FAQPage with 6–8 FAQs (valid JSON)

Reference (rewrite in your own words; do NOT copy sentences verbatim):
\"\"\"{doc_text[:10000]}\"\"\"
""".strip()

    resp = client.responses.create(
        model="gpt-5",
        input=prompt,
        text={"verbosity": "low"}
    )
    return getattr(resp, "output_text", None) or str(resp)


def markdown_to_docx(md_text: str, out_path: str):
    """
    Very lightweight Markdown -> DOCX so you can download a .docx file.
    (Headings ## and ###, simple bullets)
    """
    doc = Document()
    for line in md_text.splitlines():
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif re.match(r"^\s*-\s+", line):
            doc.add_paragraph(re.sub(r"^\s*-\s+", "", line), style="List Bullet")
        else:
            doc.add_paragraph(line)
    doc.save(out_path)


# ---------- Routes ----------
@app.get("/health")
def health():
    return {"status": "ok"}


@app.get("/download/{name}")
def download_docx(name: str):
    path = f"/tmp/{name}"
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(
        path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=name,
    )


@app.post("/upload")
async def upload(file: UploadFile = File(...)):
    # Save to a temp file
    with tempfile.NamedTemporaryFile(delete=False) as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name

    # Extract text based on extension
    name_lower = file.filename.lower()
    if name_lower.endswith(".pdf"):
        text = extract_text_from_pdf(tmp_path)
    elif name_lower.endswith(".docx"):
        text = extract_text_from_docx(tmp_path)
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type. Use .pdf or .docx")

    if not text.strip():
        raise HTTPException(status_code=422, detail="No readable text found in the document.")

    # (1) Get structured EN keywords & intents
    seo_data = generate_keywords_from_doc(text)

    # (2) Build candidate list (seed + some long-tail)
    seeds = list(dict.fromkeys(seo_data.get("seed_keywords", [])))[:15]
    longtails = list(dict.fromkeys(seo_data.get("long_tail_keywords", [])))[:15]
    candidates = seeds + longtails

    # (3) Google Trends popularity (last 3 months)
    popularity = trends_popularity(candidates, timeframe="today 3-m")

    # (4) Sort keywords by popularity (None -> bottom)
    def sort_key(k):
        v = popularity.get(k)
        return (-v) if isinstance(v, (int, float)) else float("-inf")

    ranked = sorted(candidates, key=sort_key, reverse=True)
    keywords_ranked = [{"keyword": k, "popularity": popularity.get(k)} for k in ranked]

    # (5) Compose English SEO copy that follows your template
    try:
        seo_copy_md = compose_seo_copy_en(text, ranked)
    except Exception as e:
        seo_copy_md = f"Error generating SEO copy: {e}"

    # (6) Optional DOCX export
    docx_name = None
    try:
        docx_name = f"seo-draft-{uuid.uuid4().hex}.docx"
        markdown_to_docx(seo_copy_md, f"/tmp/{docx_name}")
    except Exception:
        docx_name = None  # ignore conversion errors

    return {
        "language_detected": seo_data.get("language_detected"),
        "primary_topics": seo_data.get("primary_topics", []),
        "by_intent": seo_data.get("by_intent", {}),
        "questions": seo_data.get("questions", []),
        "keywords_ranked": keywords_ranked,
        "seed_keywords": seeds,
        "long_tail_keywords": longtails,
        "seo_copy_markdown": seo_copy_md,
        "docx_filename": docx_name,
    }


# ---------- Run (optional) ----------
if __name__ == "__main__":
    import uvicorn
    # Ensure 'python-multipart' is installed for file uploads
    # pip install python-multipart
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
